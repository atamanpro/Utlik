import os
import sqlite3
from typing import Optional, Dict, Any, List
from common_utils.s3.utils import S3yandex
from fnmatch import fnmatchcase
import json
from typing import List, Tuple
import docx
from docx.oxml.ns import qn
from docx import Document as DocxDocument


class Document:
    def __init__(self, source: str, page_content: str, metadata: Optional[Dict[str, Any]] = None):
        self.source = source
        self.page_content = page_content
        self.metadata = metadata if metadata is not None else {'source': source}


def get_uploaded_filenames(source) -> List[str]:
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute(f"SELECT filename FROM uploaded_docs WHERE global_source = ?", (source,))
        rows = cursor.fetchall()
    filenames = [row[0] for row in rows]
    return filenames


def add_filename_to_metadata(source, filename):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute(f'''INSERT INTO uploaded_docs (global_source, filename) values ('{source}', '{filename}') ; ''')


def delete_filename_from_metadata(source, filename):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute(f'''DELETE from uploaded_docs where global_source = '{source}' and filename ='{filename}' ; ''')


def load_docx_new(source, bucket: str) -> List[Document]:
    prefix = f'{current_user}/docx/'
    suffix = '.docx'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    docs = []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = obj['Body'].read()

                # Используем BytesIO для чтения содержимого файла как бинарного потока
                doc_stream = BytesIO(content)
                doc = DocxDocument(doc_stream)

                # Извлекаем текст из документа docx
                full_text = []
                image_counter = 1
                for para in doc.paragraphs:
                    # Обработка параграфов для создания ссылок на изображения
                    para_text = para.text
                    for run in para.runs:
                        for drawing in run.element.findall('.//a:blip', namespaces={
                            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                            image_rId = drawing.get(qn('r:embed'))
                            image_part = doc.part.related_parts[image_rId]
                            image_filename = f'image_{image_counter:02d}.{image_part.content_type.split("/")[-1]}'
                            image_counter += 1
                            s3_image_url = f"https://storage.yandexcloud.net/{bucket}/A100/images/{image_filename}"
                            para_text += f'\n{s3_image_url}'
                    full_text.append(para_text)
                content = '\n'.join(full_text)

                docs.append(Document(source=file, page_content=content))
            except Exception as e:
                print(f"Error reading docx file {file}: {e}")

    return docs if docs else None


def load_txts(source, bucket: str) -> List[Document]:
    prefix = f'{current_user}/txt/'
    suffix = '.txt'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    docs = []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = obj['Body'].read().decode('utf-8')
                docs.append(Document(source=file, page_content=content))
            except Exception as e:
                print(f"Error reading txt file {file}: {e}")

    return docs if docs else None


def load_jsons(source, bucket: str) -> Tuple[List[Document], List[dict]]:
    prefix = f'{current_user}/json/'
    suffix = '.json'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    json_docs, json_metadata = [], []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = json.loads(obj['Body'].read().decode('utf-8'))
                json_docs.append(content)
                json_metadata.append({'source': file})
            except Exception as e:
                print(f"Error reading json file {file}: {e}")

    return (json_docs, json_metadata) if json_docs else (None, None)


def load_documents(global_source, bucket: str, file_types: List[str]) -> dict:
    """
    Загружаем документы в зависимости от типа документа из Yandex S3
    """
    all_docs = {'txt': None, 'json': None, 'json_metadata': None, 'docx': None}
    if 'txt' in file_types:
        txt_docs = load_txts(global_source, bucket)
        all_docs['txt'] = txt_docs
    if 'json' in file_types:
        json_docs, json_metadata = load_jsons(global_source, bucket)
        all_docs['json'] = json_docs
        all_docs['json_metadata'] = json_metadata
    if 'docx' in file_types:
        docx_docs = load_docx_new(global_source, bucket)
        all_docs['docx'] = docx_docs
    return all_docs